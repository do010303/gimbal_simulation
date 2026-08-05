#!/usr/bin/env python3
"""Sinh 4 file docs/DIB_*.docx từ nội dung đã có sẵn trong docs/ (m4.md,
TEST_PLAN_RESULTS.md, docx nguồn của team). Chạy lại được khi số liệu đổi —
không sửa tay các file .docx sinh ra.

    python3 docs/tools/gen_docs.py

Định dạng bám theo docs/Tai_Lieu_Dac_Ta_KT_Milestone2_v3.docx (file mẫu tham
chiếu của team): Heading 1 cho mục lớn đánh số "1.", Heading 2 cho mục con,
bảng REQ 5-6 cột tách "bậc cao" / "đặc thù", DES gắn nhãn "-> Đáp ứng: REQ_...".
"""
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "docs"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def new_doc(title, subtitle_lines):
    d = docx.Document()
    d.add_heading(title, level=0)
    for line in subtitle_lines:
        p = d.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return d


def h1(d, text):
    return d.add_heading(text, level=1)


def h2(d, text):
    return d.add_heading(text, level=2)


def h3(d, text):
    return d.add_heading(text, level=3)


def para(d, text):
    return d.add_paragraph(text)


def bullets(d, items):
    for it in items:
        d.add_paragraph(it, style="List Bullet")


def table(d, header, rows):
    t = d.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = htext
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ---------------------------------------------------------------------------
# Shared data: hệ thống REQ bậc cao (nguồn: Drone_In_A_Box_Design_Milestones docx, muc 1)
# ---------------------------------------------------------------------------

SYSTEM_REQS = [
    ("REQ_BOX_PHY_0005", "Thời gian đóng-mở cơ cấu cửa Box",
     "Đo trong M2/M3: cửa MỞ 2.91s, ĐÓNG 3.28s."),
    ("REQ_BOX_PHY_0006", "Thời gian đóng-mở cơ cấu căn chỉnh UAV (kẹp)",
     "Đo trong M2/M3: kẹp H đóng 2.67s (19->200mm), kẹp V đóng 2.88s (14->200mm)."),
    ("REQ_BOX_FEA_0003", "Cung cấp thông tin môi trường/vị trí/trạng thái Box",
     "Message BOX_TELEMETRY trên /box_id/box_telemetry (BoxState, BoxPower, "
     "BoxEnvironment, DroneCharge)."),
    ("REQ_UAV_TALA_0007", "Sai lệch góc hạ cánh chính xác", "< 10 độ."),
    ("REQ_UAV_TALA_0008", "Sai lệch vị trí hạ cánh chính xác", "< 15 cm."),
    ("REQ_UAV_FLY_0020", "Hạ cánh vào box (kèm hạ cánh dự phòng)",
     "Drone tự động phối hợp với box để hạ cánh; khi điều kiện không an toàn "
     "(không GPS RTK / không nhận vị trí hạ cánh / box không sẵn sàng / mất "
     "tín hiệu box) thực hiện hạ cánh dự phòng (Fallback landing)."),
]

# (REQ ID, Tên yêu cầu, Mô tả (What), Tiêu chí nghiệm thu, Tham chiếu, Ghi chú)
# Nhóm theo nhu cầu chức năng thực tế, KHÔNG theo milestone — một REQ có thể
# được kiểm chứng lại ở nhiều milestone liên tiếp (liệt kê trong Tham chiếu).
FUNCTIONAL_REQS = [
    ("REQ_DIB_BOX_STATE_001", "Box phải mở đường tiếp nhận khi được yêu cầu hạ cánh",
     "Khi nhận yêu cầu hạ cánh hợp lệ lúc đang trống, box phải mở nắp/kẹp "
     "và chuyển sang trạng thái sẵn sàng đón drone trong thời gian hữu hạn.",
     "Đúng trình tự EMPTY -> IDLE -> PREPARING_FOR_LANDING -> "
     "WAITING_FOR_LANDING(7); PASS driver headless (M1, exit code 0), lặp "
     "lại đúng qua Gazebo thật (M2 test 6a, M3 PASS 8/8) và qua "
     "split-domain (M4 PASS 8/8).",
     "DES-01, DES-03, DES-06, DES-07 · M1, M2, M3, M4", ""),
    ("REQ_DIB_BOX_STATE_002", "Box phải xác nhận drone đã hạ cánh thật trước khi bảo vệ",
     "Box chỉ được kẹp giữ/đóng nắp sau khi có xác nhận thật drone đã chạm "
     "đất, không suy đoán bằng thời gian chờ.",
     "SECURING_DRONE(8) kích đúng thời điểm MAVROS landed_state=ON_GROUND "
     "(M2 test 6b); lặp lại đúng ở M3 (PASS 8/8), M4 split-domain (PASS 8/8).",
     "DES-02 · M2, M3, M4", ""),
    ("REQ_DIB_BOX_STATE_003", "Box phải tự khép chu trình bảo vệ và chuyển sang sạc",
     "Sau khi bảo vệ drone xong (kẹp, đóng nắp, cắt điện an toàn), box phải "
     "tự động chuyển sang trạng thái sạc, không cần can thiệp tay.",
     "Chuỗi SECURING_DRONE(8) -> CHARGING(9) hoàn tất tự động; đo M4: "
     "POWER_OFF -> CHARGING ~4.6s sau khi cắt điện.",
     "DES-01, DES-08 · M1, M2, M3, M4", ""),
    ("REQ_DIB_BOX_STATE_004", "Trạng thái cơ cấu chấp hành báo cáo phải đúng thực tế vật lý",
     "Trạng thái cửa/kẹp mà box báo cho hệ thống phải phản ánh đúng chuyển "
     "động vật lý thật, không phải giá trị giả định — tránh hệ thống "
     "'tưởng' cửa đã mở khi thực tế còn kẹt.",
     "LidStatus/ClampStatus suy từ /joint_states thật; đo M2: lid chuyển "
     "động >0.5 rad, clamp đạt ~200mm trên Gazebo thật.",
     "DES-02 · M2", "Liên quan REQ_BOX_PHY_0005/0006 (thời gian đóng-mở "
     "đã đo ở mức hệ thống)."),
    ("REQ_DIB_COORD_001", "Drone không được hạ cánh khi box chưa sẵn sàng đón",
     "Drone phải chờ box xác nhận đã sẵn sàng (mở nắp) trước khi bắt đầu "
     "quy trình hạ cánh thị giác — tránh hạ vào box còn đóng.",
     "Controller chỉ rời trạng thái chờ sang hạ cánh SAU KHI box báo "
     "WAITING_FOR_LANDING(7); xác nhận bằng nhân quả qua monitor thụ động, "
     "PASS 8/8 (M3), PASS 8/8 (M4 split-domain).",
     "DES-03, DES-04 · M3, M4", ""),
    ("REQ_DIB_COORD_002", "Một lượt bay đầy đủ phải khép kín tới sạc với sai số đo được",
     "Vòng đời từ cất cánh tới box sạc xong phải chạy trọn, và sai số hạ "
     "cánh phải đo định lượng được (không chỉ quan sát định tính).",
     "PASS 8/8 (M3): sai số vị trí 4.9cm/1.6cm, yaw 0.00 độ; PASS 8/8 (M4, "
     "split-domain): sai số 2.2cm, yaw 0.01 độ.",
     "DES-04, DES-05 · M3, M4", "Đạt REQ_UAV_TALA_0007/0008 với biên độ lớn."),
    ("REQ_DIB_NET_001", "Hệ thống phải hoạt động đúng khi box và drone tách hai domain mạng",
     "Toàn bộ hành vi phối hợp box-drone (các REQ ở trên) phải giữ nguyên "
     "khi box và drone chạy trên hai ROS_DOMAIN_ID tách biệt — bước chuẩn "
     "bị triển khai trên hai máy vật lý qua LAN thật.",
     "PASS 8/8 driver tự động, split-domain (M4); bằng chứng dữ liệu thật "
     "sự qua cầu bằng đo cổng UDP (domain 0 -> 7400, domain 42 -> 17900).",
     "DES-06, DES-07, DES-08, DES-09 · M4", ""),
    ("REQ_DIB_NET_002", "Mất kết nối box-drone giữa chừng phải dẫn tới hạ cánh an toàn trong thời hạn xác định",
     "Nếu đường truyền giữa box và drone bị gián đoạn trong lúc đang bắt "
     "tay, drone không được treo vô thời hạn mà phải chuyển sang hạ cánh "
     "dự phòng trong thời hạn xác định trước.",
     "Kịch bản tắt cầu giữa chuyến bay đang sống (M4): box đứng nguyên "
     "EMPTY suốt 39.5s, đúng 30.0s FALLBACK kích hoạt tự nhiên, hệ thống "
     "chuyển AUTO.LAND an toàn.",
     "DES-04 · M4", "Hiện thực REQ_UAV_FLY_0020 (hạ cánh dự phòng) trong "
     "điều kiện mất mạng, không chỉ mất GPS/box."),
    ("REQ_DIB_NET_003", "Người vận hành phải phát hiện được khi mất kết nối box-drone",
     "Việc mất cầu kết nối phải quan sát được trực tiếp, không chỉ suy ra "
     "gián tiếp sau khi hệ thống timeout — rút ngắn thời gian chẩn đoán "
     "sự cố.",
     "Tín hiệu heartbeat định kỳ 10s từ cầu; mất tín hiệu >10s = cầu chết "
     "(kiểm bằng smoke test 13s, M5).",
     "DES-11 · M5", ""),
    ("REQ_DIB_NET_004", "Box phải phát hiện khi dữ liệu trạng thái drone ngừng cập nhật",
     "Trong lúc chờ drone hạ cánh, box phải nhận biết được khi không còn "
     "nhận telemetry mới từ drone quá lâu, tránh dựa vào dữ liệu cũ mà "
     "không hay biết.",
     "Cảnh báo khi telemetry cũ quá 5.0s trong lúc WAITING_FOR_LANDING; "
     "không đổi hành vi FSM (kiểm bằng build + test thủ công, M5).",
     "DES-12 · M5", ""),
    ("REQ_DIB_OPS_001", "Người vận hành phải kiểm tra được điều kiện bay đủ trong một bước",
     "Trước mỗi lượt bay, phải có cách kiểm nhanh, không phụ thuộc trí nhớ "
     "nhiều bước tay, để xác nhận hệ thống đủ điều kiện bay.",
     "go/no-go chạy đúng ở cả chế độ có box và không có box, báo đúng "
     "GO/NO-GO theo trạng thái pipeline thật (M5).",
     "DES-10 · M5", ""),
    ("REQ_DIB_OPS_002", "Người mới phải chạy được hệ thống chỉ từ tài liệu",
     "Toàn bộ hướng dẫn cài đặt, chạy và xử lý sự cố phải nằm ở một nơi tra "
     "cứu được, đúng đường dẫn, không phụ thuộc tài liệu lỗi thời.",
     "README là mục lục chạy-được duy nhất, 0 dấu vết tài liệu cũ "
     "(gimbal_simulation) trong repo; rà theo góc nhìn máy trống xác nhận "
     "đủ bước (M5).",
     "— · M5", "DoD Milestone M5 (mục 3, kế hoạch gốc)."),
    ("REQ_DIB_OPS_003", "Sản phẩm bàn giao không được lẫn rác/bản sao cũ",
     "Cây thư mục đóng gói và máy phát triển không được chứa bản sao/rác "
     "gây nhầm lẫn cho người triển khai lại.",
     "Đã xoá phần rác xác nhận an toàn (trùng lặp, không .git); các phần "
     "còn việc dở của người khác đã xác định rõ, không đưa vào sản phẩm "
     "giao (M5).",
     "— · M5", ""),
]

DES_LIST = [
    ("DES-01", "Contract-first FSM (M1)",
     "box_state_manager điều khiển thuần bằng contract dib_msgs (BoxCmd "
     "vào, BoxTelemetry/BoxState ra) — kích thích/quan sát được bằng driver "
     "script thuần Python, không cần Gazebo hay drone stack.",
     "REQ_DIB_BOX_STATE_001, REQ_DIB_BOX_STATE_003"),
    ("DES-02", "Mẫu Adapter, không sửa hai đầu (M2)",
     "Chèn lớp box_hardware_adapter dịch dib_msgs <-> JointTrajectory giữa "
     "box_manager và box_simulation, thêm component thứ tư "
     "mavros_to_dib_telemetry cho hướng drone->box — không sửa mã nguồn "
     "hai hệ đã có.",
     "REQ_DIB_BOX_STATE_002, REQ_DIB_BOX_STATE_004"),
    ("DES-03", "Module BoxLink + 2 state bắt tay (M3)",
     "Toàn bộ logic bắt tay box (gửi REQUEST_LANDING, chờ WAITING_FOR_LANDING) "
     "đóng gói trong module BoxLink riêng (controller chỉ giữ một instance "
     "box_link_), hai trạng thái PRELANDING_CHECK/WAIT_BOX_READY thêm vào "
     "FSM offboard_precland_controller.",
     "REQ_DIB_BOX_STATE_001, REQ_DIB_COORD_001"),
    ("DES-04", "Nghiệm thu bằng monitor thụ động + tiêu chí nhân quả (M3, M4)",
     "Driver nghiệm thu chỉ SUBSCRIBE (không publish/gọi service), chấm PASS "
     "dựa trên thứ tự thời gian giữa hai FSM (box rời EMPTY sau khi drone "
     "vào WAIT_BOX_READY; drone vào START sau khi box đạt "
     "WAITING_FOR_LANDING) — loại trừ khả năng PASS giả do harness tự kích "
     "thích hệ thống.",
     "REQ_DIB_COORD_001, REQ_DIB_COORD_002, REQ_DIB_NET_002"),
    ("DES-05", "Hợp nhất world + marker trên thân box (M3)",
     "Marker fractal đặt trực tiếp trên sàn box trong cùng world Gazebo với "
     "drone, thay vì hai world tách rời — cho phép vòng kín thị giác thật "
     "chạy tới CHARGING.",
     "REQ_DIB_COORD_002"),
    ("DES-06", "DDS-Router 2.2.0 + whitelist-interfaces per-participant (M4)",
     "Cầu domain mặc định là DDS-Router 2.2.0 (Fast DDS 2.14.0, cùng dòng "
     "2.x với rmw_fastrtps Humble). Mỗi participant trong config khai "
     "whitelist-interfaces: [\"127.0.0.1\"] (tag per-participant, không "
     "phải top-level) để tránh locator LAN thật bị loại. allowlist khai "
     "tường minh cả name lẫn type DDS-mangled cho message tuỳ biến.",
     "REQ_DIB_BOX_STATE_001, REQ_DIB_NET_001"),
    ("DES-07", "Lệnh drone->box chuyển từ service sang topic (M4)",
     "b2/cmd (service) không còn dùng cho đường bắc cầu vì DDS-Router 2.2.0"
     " không bao giờ route reply qua domain (bug thật, đo với cả message "
     "tuỳ biến lẫn service chuẩn ROS 2). Lệnh REQUEST_LANDING/TURN_OFF_DRONE "
     "chuyển sang topic b2/drone_cmd (dib_msgs/BoxCmd, thêm field "
     "command/reserve/agent_id). Service b2/cmd vẫn giữ cho vai trò "
     "operator/server nội bộ, chỉ không đi qua ranh giới domain nữa.",
     "REQ_DIB_BOX_STATE_001, REQ_DIB_NET_001"),
    ("DES-08", "Fixture SITL /dock/drone_power thay đường điện vật lý (M4)",
     "Trên phần cứng thật, box cắt điện làm máy tính drone tắt hẳn nên "
     "d1/telemetry im miễn phí. Trong SITL MAVROS vẫn chạy, nên "
     "box_hardware_adapter publish topic /dock/drone_power "
     "(RELIABLE/TRANSIENT_LOCAL) và mavros_to_dib_telemetry ngừng phát khi "
     "nhận false — topic này phải nằm trong allowlist bắc cầu, thiếu nó box "
     "kẹt vĩnh viễn ở SECURING_DRONE.",
     "REQ_DIB_BOX_STATE_003, REQ_DIB_NET_001"),
    ("DES-09", "Cầu dự phòng dib_domain_bridge (M4)",
     "Gói ros-humble-domain-bridge (cài từ apt, vài giây) làm cầu dự phòng "
     "khi máy không build được DDS-Router 2.2.0 từ nguồn (~4 phút). Bắc cầu "
     "cùng 3 giao diện hợp đồng + fixture /dock/drone_power, cộng heartbeat "
     "log định kỳ (DES-11).",
     "REQ_DIB_NET_001 (giảm rủi ro triển khai — không bắt buộc bởi REQ nào "
     "trực tiếp)"),
    ("DES-10", "Cổng go/no-go một lệnh (M5)",
     "scripts/go_no_go.sh gộp verify_build_env.sh (môi trường build), kiểm "
     "tiến trình cũ còn sót (cùng pattern với stop_pipeline.sh), và 3 mục "
     "kiểm trước bay (use_sim_time, controller active, MAVROS connected) — "
     "tự nhận biết chế độ có/không có box (README mục 2 và 3).",
     "REQ_DIB_OPS_001"),
    ("DES-11", "Heartbeat log định kỳ trong dib_split_bridge (M5)",
     "Timer 10s log uptime + domain đang bắc cầu trong "
     "dib_split_bridge.cpp — biến 'cầu chết im lặng' (DDS-Router và "
     "domain_bridge đều không log từng message) thành 'thiếu dòng heartbeat "
     "gần nhất là cầu chết', quan sát được từ ngoài mà không cần công cụ đo "
     "mạng.",
     "REQ_DIB_NET_003"),
    ("DES-12", "Cảnh báo throttle telemetry cũ trong box_state_manager (M5)",
     "Lưu last_drone_telemetry_time_ (cùng mẫu last_*_time_ đã dùng ở "
     "offboard_precland_controller), so sánh ngưỡng 5.0s, RCLCPP_WARN_THROTTLE"
     " khi đang WAITING_FOR_LANDING — chỉ chẩn đoán, không đổi hành vi FSM "
     "(SECURING_DRONE/POWER_OFF cố ý không kiểm vì staleness ở đó là trigger "
     "CHARGING bình thường, không phải lỗi).",
     "REQ_DIB_NET_004"),
]


# ---------------------------------------------------------------------------
# DOC 1 — Mô tả công việc của các Milestone
# ---------------------------------------------------------------------------

def gen_doc1():
    d = new_doc(
        "MÔ TẢ CÔNG VIỆC CỦA CÁC MILESTONE",
        ["Hệ thống Drone-in-a-Box — PX4 SITL Precision Landing + Box Manager + DDS-Router",
         "Sinh tự động từ docs/tools/gen_docs.py — nguồn: m4.md, TEST_PLAN_RESULTS.md, "
         "Drone_In_A_Box_Design_Milestones docx gốc"],
    )

    h1(d, "1. M1 — Build & chạy độc lập box_manager")
    para(d, "Trạng thái: Đã hoàn thành (20/07/2026)")
    h2(d, "Mô tả")
    para(d, "Vấn đề cần giải quyết: box_manager có đủ source C++ nhưng chưa "
            "được build/chạy trong SITL — README của nó trỏ sang một kho "
            "dib_msgs GitLab riêng, không trỏ bộ dib_msgs local. Chưa có "
            "bằng chứng nào cho thấy state machine của box chạy đúng đặc tả.")
    para(d, "Yêu cầu liên quan: REQ_DIB_BOX_STATE_001, REQ_DIB_BOX_STATE_003 "
            "(kiểm bằng driver headless ở mức này — chưa tích hợp phần "
            "cứng/FSM drone thật) — xem docs/DIB_REQ_DES.docx.")
    para(d, "Thiết kế: DES-01 (contract-first FSM) — xem docs/DIB_REQ_DES.docx.")
    para(d, "DoD: colcon build sạch; log chuyển trạng thái đi đúng enum khi "
            "kích lệnh tay qua driver.")
    h2(d, "Hướng giải quyết vấn đề (insight)")
    bullets(d, [
        "Nút thắt không phải ở logic mà ở dependency: chỉ cần trỏ box_manager "
        "sang dib_msgs local là build và chạy được — logic FSM vốn đã đúng.",
        "Vì FSM chỉ phụ thuộc contract, viết driver headless m1_state_test.py "
        "(giả BoxCmd + một mẫu DroneTelemetry) đủ để nghiệm thu TRƯỚC khi có "
        "Gazebo — cô lập được rủi ro.",
        "FSM chạy đúng dù box_info.latitude/longitude = 0 -> GPS không phải "
        "phụ thuộc của luồng landing (xác nhận lại ở M2 known-gap navsat).",
    ])
    h2(d, "Kết quả đạt được")
    para(d, "Driver theo dõi /box/state đi trọn tới CHARGING -> RESULT: PASS "
            "(exit code 0). REQ_DIB_BOX_STATE_001, REQ_DIB_BOX_STATE_003 đạt "
            "ở mức driver headless. Đây là nền để M2/M3 tích hợp phần thật.")

    h1(d, "2. M2 — Box hardware adapter + telemetry bridge")
    para(d, "Trạng thái: Đã hoàn thành (26/07/2026)")
    h2(d, "Mô tả")
    para(d, "Vấn đề cần giải quyết: Hai hệ nói hai ngôn ngữ khác nhau: "
            "box_simulation điều khiển bằng JointTrajectory thô của "
            "ros2_control, còn box_manager nói dib_msgs (LidCmd/ClampCmd/...)."
            " box_manager cũng cần biết drone đã hạ thật để chuyển "
            "SECURING_DRONE — tín hiệu này phải đến từ PX4, không phải "
            "phỏng đoán.")
    para(d, "Yêu cầu liên quan: REQ_DIB_BOX_STATE_002, REQ_DIB_BOX_STATE_004 "
            "— xem docs/DIB_REQ_DES.docx.")
    para(d, "Thiết kế: DES-02 (mẫu Adapter) — xem docs/DIB_REQ_DES.docx.")
    para(d, "DoD: gọi /lid/cmd OPEN -> lid Gazebo mở thật -> box_manager tự "
            "chuyển PREPARING_FOR_LANDING -> WAITING_FOR_LANDING không cần "
            "publish tay; box_manager tự chuyển sang SECURING_DRONE khi PX4 "
            "land detector báo landed thật.")
    h2(d, "Hướng giải quyết vấn đề (insight)")
    bullets(d, [
        "Tín hiệu 'đã hạ' PHẢI lấy từ extended_state.landed_state của PX4 "
        "(land detector thật), không phải timer — điều khiến bước securing "
        "kích đúng thời điểm.",
        "Khác biệt đo lường: 29 'mismatch' thoáng qua là độ trễ transient "
        "của passthrough, không phải lỗi; DoD thực chất đạt ngay — bài học "
        "về việc chọn đúng metric.",
        "Hạ tầng bắt buộc (làm một lần): build gz_ros2_control cho Harmonic "
        "(bản apt là Fortress -> gz server segfault khi spawn box) và đổi "
        "tên plugin ign_ros2_control -> gz_ros2_control trong box.xacro.",
        "Known gap không chặn M2: cảm biến navsat trong box.xacro chưa có "
        "ros_gz_bridge sang NavSatFix, nên box lat/lon = 0 trong SITL — "
        "không ảnh hưởng landing (đã chứng minh ở M1).",
    ])
    h2(d, "Kết quả đạt được")
    para(d, "3/3 test PASS trên hệ thống thật: 6a-unit (logic adapter), "
            "6a-full (lid chuyển động >0.5 rad, clamp đạt ~200mm trên Gazebo "
            "thật), 6b (MAVROS landed passthrough đúng thời điểm). "
            "REQ_DIB_BOX_STATE_002, REQ_DIB_BOX_STATE_004 đạt.")

    h1(d, "3. M3 — Bắt tay FSM drone<->box qua BoxLink")
    para(d, "Trạng thái: Đã hoàn thành (30/07/2026)")
    h2(d, "Mô tả")
    para(d, "Vấn đề cần giải quyết: Ở bản cũ, GOTO_BOX nhảy thẳng sang START "
            "chỉ dựa trên bán kính GPS: không gọi BoxCmd, không chờ box mở "
            "nắp, không bắt tay. Cần hai FSM (drone và box) phối hợp qua "
            "contract, và phải chứng minh sự khép vòng là NHÂN QUẢ (do "
            "REQUEST_LANDING gây ra) chứ không phải sản phẩm phụ của trigger "
            "tay còn sót.")
    para(d, "Yêu cầu liên quan: REQ_DIB_COORD_001, REQ_DIB_COORD_002 (tái "
            "xác nhận REQ_DIB_BOX_STATE_001/002/003 với vòng lặp thật) — "
            "xem docs/DIB_REQ_DES.docx.")
    para(d, "Thiết kế: DES-03 (module BoxLink), DES-04 (monitor thụ động), "
            "DES-05 (world hợp nhất) — xem docs/DIB_REQ_DES.docx.")
    para(d, "DoD: sitl_precland.launch.py chạy, /lander/state đi đủ chuỗi "
            "state mới; hệ thống không còn phụ thuộc stub sim_box_manager.py "
            "(đã loại bỏ).")
    h2(d, "Hướng giải quyết vấn đề (insight)")
    bullets(d, [
        "Hai tiêu chí nhân quả then chốt: box phải rời EMPTY SAU khi drone "
        "vào WAIT_BOX_READY, và drone phải vào START SAU khi box đạt "
        "WAITING_FOR_LANDING(7). Đúng thứ tự này mới chứng minh chính "
        "REQUEST_LANDING gây ra chuyển động, loại trừ trigger tay còn sót.",
        "Monitor thụ động (chỉ nghe topic, không publish) đảm bảo PASS "
        "không phải sản phẩm phụ của việc harness tự kích thích hệ thống.",
        "Đặt bắt tay vào BoxLink giữ file controller ~1500 dòng không bị "
        "nhiễm logic service/MAVLink; đúng hướng layer hoá đã tự vạch trong "
        "docs/refactoring_architecture.md.",
    ])
    h2(d, "Kết quả đạt được")
    para(d, "PASS 8/8 trên monitor thụ động (gồm hai tiêu chí nhân quả), "
            "vòng kín chạy tới CHARGING(9). Sai số hạ cánh THẬT đo được lần "
            "đầu = 4.9 cm. Các mốc con: M3a bắt tay FSM, M3b world hợp nhất "
            "(PASS 7/7), M3c khép vòng tới CHARGING (PASS 8/8), M3d dọn log "
            "+ báo vị trí. REQ_DIB_COORD_001, REQ_DIB_COORD_002 đạt.")

    h1(d, "4. M4 — Tách domain box<->drone, bắc cầu hợp đồng (DDS-Router)")
    para(d, "Trạng thái: Đã hoàn thành (05/08/2026)")
    h2(d, "Mô tả")
    para(d, "Vấn đề cần giải quyết: M1-M3 chạy toàn bộ trong một "
            "ROS_DOMAIN_ID (single-domain). Kiến trúc mục tiêu của DIB tách "
            "box và drone thành hai máy vật lý riêng (LAN thật), nghĩa là "
            "hai ROS_DOMAIN_ID riêng — cần chứng minh trước, trên một host, "
            "rằng vòng đời khép được qua ranh giới domain mà không sửa logic "
            "nghiệp vụ hai FSM đã có (M1/M3), chỉ thêm hạ tầng bắc cầu.")
    para(d, "Yêu cầu liên quan: REQ_DIB_NET_001, REQ_DIB_NET_002 (+ tái xác "
            "nhận REQ_DIB_BOX_STATE_*/REQ_DIB_COORD_* qua split-domain) — "
            "xem docs/DIB_REQ_DES.docx.")
    para(d, "Thiết kế: DES-06 (DDS-Router + whitelist-interfaces), DES-07 "
            "(lệnh drone->box chuyển service->topic), DES-08 (fixture "
            "/dock/drone_power), DES-09 (cầu dự phòng dib_domain_bridge) — "
            "xem docs/DIB_REQ_DES.docx.")
    para(d, "DoD: vòng đời EMPTY -> ... -> CHARGING chạy trọn qua hai domain "
            "tách biệt (box=42, drone=0); điểm 8/8 từ driver tự động; bằng "
            "chứng nhân quả tắt cầu giữa chừng một chuyến bay thật.")
    h2(d, "Hướng giải quyết vấn đề (insight)")
    bullets(d, [
        "DDS-Router 3.x (bản mới nhất) là ngõ cụt: Humble link Fast DDS "
        "2.6.11, mọi bản DDS-Router 3.x kéo theo Fast DDS 3.x — một "
        "participant 3.6 không discovery được endpoint 2.6, router chạy, "
        "báo 'running', nhưng bắc cầu 0 message. Đo hai lần độc lập "
        "(2026-07-30 và build lại 2026-08-05) cho kết quả giống hệt. "
        "DDS-Router 2.2.0 (Fast DDS 2.14.0, cùng dòng 2.x) mới nói chuyện "
        "được với 2.6.",
        "'ROS_LOCALHOST_ONLY=0 mới bắc cầu được' là kết luận sai (2026-07-30)."
        " Nguyên nhân thật: whitelist-interfaces là tag per-participant, "
        "không phải top-level — đặt sai chỗ thì node ROS âm thầm loại "
        "locator LAN thật mà DDS-Router quảng bá.",
        "Service /b2/cmd không bao giờ route được reply qua domain — bug "
        "thật của DDS-Router 2.2.0, đo với cả message tuỳ biến lẫn service "
        "chuẩn ROS 2 (demo_nodes_cpp AddTwoInts), không sửa được bằng cấu "
        "hình. Giải pháp: chuyển lệnh drone->box sang topic b2/drone_cmd — "
        "an toàn vì reply cũ vốn vô nghĩa (box_state_manager set "
        "success=true trước khi xử lý gì, telemetry mới là xác nhận thật).",
        "Lần chạy Gazebo đầu tiên (E.2): mọi giao diện hợp đồng qua được "
        "nhưng box kẹt vĩnh viễn ở SECURING_DRONE — nguyên nhân là thiếu "
        "/dock/drone_power (fixture SITL) trong allowlist, không phải lỗi "
        "cầu. Dấu hiệu xác nhận: không có dòng 'Dock power OFF: stopping "
        "publishing d1/telemetry' bên drone.",
    ])
    h2(d, "Kết quả đạt được")
    para(d, "PASS 8/8 chính thức bằng driver tự động (không đọc log tay), "
            "split-domain, gồm 2 tiêu chí nhân quả. Sai số bám hạ cánh "
            "(aim_error) lần PASS cuối = 2.2 cm, sai số yaw = 0.01°. "
            "POWER_OFF -> CHARGING sau khi cắt điện ~4.6s (khớp ngưỡng thiết "
            "kế 5s). Bằng chứng cầu thật sự mang dữ liệu qua cổng UDP: "
            "domain 0 -> 7400, domain 42 -> 17900, ddsrouter là tiến trình "
            "duy nhất giữ cả hai cổng. Thêm bằng chứng nhân quả kịch bản 3: "
            "tắt cầu giữa chừng một chuyến bay đang sống -> box đứng nguyên "
            "ở EMPTY suốt 39.5s, đúng 30.0s FALLBACK in ra tự nhiên, hệ "
            "thống rơi về AUTO.LAND an toàn (aim_error=2.385m, kém chính xác "
            "hơn ~100 lần so với hạ cánh thị giác — đúng hành vi thiết kế, "
            "không phải lỗi). REQ_DIB_NET_001, REQ_DIB_NET_002 đạt.")

    h1(d, "5. M5 — Hardening + debug + đóng gói")
    para(d, "Trạng thái: Đang triển khai (kế hoạch/định hướng — cập nhật khi "
            "hoàn thành)")
    h2(d, "Mô tả")
    para(d, "Vấn đề cần giải quyết: M1-M4 đã chứng minh logic đúng nhưng "
            "chưa đóng gói được cho người ngoài dùng lại: chưa có cổng "
            "go/no-go một lệnh trước mỗi lượt bay, runbook còn phân mảnh "
            "(một phần trong README, một phần trong docs/m4_split_domain_test),"
            " cầu split-domain (M4) im lặng khi chết (không log từng "
            "message), box không phát hiện telemetry drone quá cũ, và máy "
            "phát triển còn rác ngoài repo (bản sao cũ box_manager/... không "
            "còn cần).")
    para(d, "Yêu cầu liên quan: REQ_DIB_NET_003, REQ_DIB_NET_004, "
            "REQ_DIB_OPS_001, REQ_DIB_OPS_002, REQ_DIB_OPS_003 — xem "
            "docs/DIB_REQ_DES.docx.")
    para(d, "Thiết kế dự kiến: DES-10 (cổng go/no-go), DES-11 (heartbeat "
            "cầu), DES-12 (cảnh báo telemetry cũ) — xem docs/DIB_REQ_DES.docx.")
    para(d, "DoD (theo bảng milestone gốc): luồng SITL chạy ổn; README "
            "hướng dẫn chạy hệ thống đầy đủ; deadline đóng gói sản phẩm "
            "10/08/2026.")
    h2(d, "Hướng giải quyết vấn đề (định hướng — sẽ cập nhật thành DES khi "
            "hoàn thành và đã thử-sai)")
    bullets(d, [
        "Go/no-go: gộp verify_build_env.sh (môi trường build) + kiểm tiến "
        "trình cũ còn sót + 3 mục kiểm trước bay đã có sẵn rải rác trong "
        "README, thành một script duy nhất, tự nhận biết chế độ có/không "
        "có box.",
        "Runbook một file: README.md giữ vai trò mục lục chạy-được duy "
        "nhất; Phụ lục A.10 trỏ sang docs/m4_split_domain_test/README.md "
        "cho chi tiết split-domain thay vì trùng lặp nội dung.",
        "Hardening tối thiểu, không đụng đường bay chính: heartbeat log "
        "định kỳ cho cầu split-domain (vấn đề gặp thật ở kịch bản 3 của "
        "M4 — cầu chết im lặng); cảnh báo throttle khi drone_telemetry cũ "
        "quá ngưỡng trong box_state_manager (theo đúng mẫu last_*_time_ đã "
        "dùng ở offboard_precland_controller) — chỉ chẩn đoán, không đổi "
        "hành vi FSM.",
        "Dọn máy phát triển: xác nhận .git lồng bên trong các thư mục rác "
        "trước khi động vào — tránh mất việc chưa lưu của người khác (rút "
        "kinh nghiệm trực tiếp khi rà soát: phát hiện các thư mục ngoài "
        "repo có .git riêng với commit/thay đổi CHƯA đẩy lên remote gốc).",
    ])
    h2(d, "Kết quả đạt được (cập nhật tới thời điểm hiện tại)")
    para(d, "scripts/go_no_go.sh + docs/go_no_go.md đã có, chạy được ở cả "
            "chế độ có box (README mục 2) và không box (README mục 3). "
            "README mục 2/3 đã trỏ vào go/no-go. Heartbeat 10s đã thêm vào "
            "dib_split_bridge.cpp, cảnh báo throttle telemetry cũ đã thêm "
            "vào box_state_manager.cpp — cả hai đã build sạch. Toàn bộ tài "
            "liệu test M1-M4 đã mở khỏi .gitignore và commit vào nhánh dev "
            "chính thức. Còn lại: hoàn tất bộ 4 tài liệu tái cấu trúc này, "
            "và rà soát cuối trước hạn đóng gói 10/08/2026.")

    h1(d, "6. Đối chiếu với yêu cầu bậc hệ thống (Traceability, mục 4.4 cũ)")
    para(d, "Trạng thái đánh giá trung thực trên hệ thống SITL hiện tại: ĐẠT / "
            "ĐÃ ĐO / Đạt một phần / Chưa kiểm chứng. Kế thừa nguyên các dòng "
            "M1-M3 từ tài liệu gốc, bổ sung dòng M4 (đã nghiệm thu 05/08) và "
            "M5 (đang triển khai).")
    table(d, ["Mã yêu cầu", "Milestone", "Trạng thái", "Ghi chú / khoảng trống"], [
        ("REQ_BOX_FEA_0003", "M1, M2", "ĐẠT",
         "Telemetry đủ: GPS thật (fixture) + môi trường (giá trị mô phỏng "
         "qua box_env_publisher) + field additive (status_door/status_hold/"
         "is_empty/connected/air_conditioner/status_power)."),
        ("REQ_BOX_PHY_0005", "M2", "ĐÃ ĐO",
         "Đo trong loop M3: cửa MỞ 2.91s, ĐÓNG 3.28s."),
        ("REQ_BOX_PHY_0006", "M2", "ĐÃ ĐO",
         "Đo trong loop M3: kẹp H đóng 2.67s (19->200mm), kẹp V đóng 2.88s "
         "(14->200mm)."),
        ("REQ_UAV_TALA_0007", "M3", "ĐẠT",
         "Bay Gazebo 2026-08-03: yaw_error=0.00 độ < 10 độ."),
        ("REQ_UAV_TALA_0008", "M3", "ĐẠT",
         "Sai lệch vị trí hạ cánh 1.6cm (bay 2026-08-03) / 4.9cm (M3c) "
         "< 15cm."),
        ("REQ_UAV_FLY_0020", "M3", "ĐẠT",
         "Bay require_rtk:=true: PRELANDING_CHECK->FALLBACK đúng thiết kế, "
         "DroneTelemetry.error=[2] xác nhận hạ cánh dự phòng."),
        ("REQ_BOX_FEA_0003", "M4", "ĐẠT (bổ sung split-domain)",
         "Telemetry vẫn đúng khi bắc cầu qua 2 ROS_DOMAIN_ID tách biệt "
         "(b2/telemetry, d1/telemetry qua DDS-Router 2.2.0) — xem docs/m4.md."),
        ("REQ_UAV_FLY_0020", "M4", "ĐẠT (bổ sung split-domain)",
         "Hạ cánh + fallback vẫn đúng qua split-domain: kịch bản tắt cầu "
         "giữa một chuyến bay đang sống xác nhận FALLBACK kích đúng 30.0s, "
         "rơi về AUTO.LAND an toàn."),
        ("REQ_DIB_NET_003, REQ_DIB_NET_004, REQ_DIB_OPS_001..003", "M5",
         "Đang triển khai",
         "go/no-go + heartbeat cầu + cảnh báo telemetry cũ đã code/build; "
         "dọn rác đóng gói + rà runbook máy sạch chưa chốt hẳn."),
    ])

    d.save(f"{OUT_DIR}/DIB_M4_Muc4_Mo_Ta_Cong_Viec.docx")
    print("Đã sinh:", f"{OUT_DIR}/DIB_M4_Muc4_Mo_Ta_Cong_Viec.docx")


# ---------------------------------------------------------------------------
# DOC 2 — REQ + DES gộp (đối chiếu qua lại)
# ---------------------------------------------------------------------------

def gen_doc2():
    d = new_doc(
        "REQUIREMENTS & DESIGN DESCRIPTION",
        ["Hệ thống Drone-in-a-Box — Requirement (What) & Design (How), gộp một "
         "file để đối chiếu qua lại",
         "Sinh tự động từ docs/tools/gen_docs.py"],
    )

    h1(d, "1. Yêu cầu (REQ) — trả lời What, không mô tả cách thức")
    h2(d, "1.1. Yêu cầu bậc cao (System-Level, nguồn: đặc tả hệ thống DIB)")
    table(d, ["REQ ID", "Tên yêu cầu", "Mô tả (What)", "Giá trị/Tham chiếu"],
          SYSTEM_REQS)

    h2(d, "1.2. Yêu cầu đặc thù (theo nhu cầu chức năng, không theo Milestone)")
    para(d, "Một REQ có thể được kiểm chứng lại ở nhiều milestone liên tiếp — "
            "cột 'Tham chiếu' liệt kê đủ các milestone đó cùng mã DES (mục 2 "
            "bên dưới) hiện thực nó, phục vụ đối chiếu qua lại theo yêu cầu "
            "traceability của Appendix C, NASA Systems Engineering Handbook.")
    table(
        d,
        ["REQ ID", "Tên yêu cầu", "Mô tả (What)", "Tiêu chí nghiệm thu",
         "Tham chiếu", "Ghi chú"],
        FUNCTIONAL_REQS,
    )

    h1(d, "2. Thiết kế (DES) — trả lời How, viết sau khi đã thử-sai")
    para(d, "Mỗi DES gắn nhãn '-> Phục vụ REQ' để đối chiếu ngược lại bảng ở "
            "mục 1.2. DES-01..09 đã thực hiện và nghiệm thu (M1-M4); "
            "DES-10..12 thuộc M5, đã code+build, đang hoàn thiện nghiệm thu.")
    table(d, ["DES ID", "Tên", "Mô tả (How)", "Phục vụ REQ"], DES_LIST)

    d.save(f"{OUT_DIR}/DIB_REQ_DES.docx")
    print("Đã sinh:", f"{OUT_DIR}/DIB_REQ_DES.docx")


# ---------------------------------------------------------------------------
# DOC 3 — Report giai đoạn
# ---------------------------------------------------------------------------

def gen_doc3():
    d = new_doc(
        "REPORT GIAI ĐOẠN",
        ["Hệ thống Drone-in-a-Box — những gì đã làm trong từng Milestone và "
         "insight rút ra (kể cả các phép thử fail)",
         "Sinh tự động từ docs/tools/gen_docs.py — nguồn: TEST_PLAN_RESULTS.md, "
         "m3.md, m4.md"],
    )

    h1(d, "M1 — Build & chạy độc lập box_manager (20/07/2026)")
    para(d, "Đã làm: gỡ phụ thuộc box_manager khỏi kho dib_msgs GitLab riêng, "
            "trỏ sang dib_msgs local trong repo; viết driver headless "
            "m1_state_test.py giả lập BoxCmd + DroneTelemetry để nghiệm thu "
            "FSM mà không cần Gazebo.")
    para(d, "Insight: nút thắt tưởng là logic hoá ra chỉ là dependency path — "
            "bài học về việc xác minh giả thuyết trước khi bắt tay sửa. FSM "
            "không phụ thuộc GPS (lat/lon=0 vẫn chạy đúng), tách được rủi ro "
            "cảm biến khỏi rủi ro state machine ngay từ milestone đầu.")

    h1(d, "M2 — Box hardware adapter + telemetry bridge (26/07/2026)")
    para(d, "Đã làm: viết box_hardware_adapter (C++) dịch dib_msgs <-> "
            "JointTrajectory; viết mavros_to_dib_telemetry ánh xạ MAVROS -> "
            "d<id>/telemetry; build gz_ros2_control từ nguồn cho Harmonic.")
    para(d, "Insight (kể cả phép thử không đạt lúc đầu): ban đầu đo được "
            "29 lần 'mismatch' giữa trạng thái báo và trạng thái thật — nghi "
            "ngờ là lỗi logic adapter. Điều tra kỹ hơn cho thấy đây chỉ là "
            "độ trễ transient bình thường của việc passthrough qua "
            "/joint_states (thời gian giữa lúc gửi lệnh và lúc khớp vật lý "
            "di chuyển xong), không phải lỗi — DoD thực chất đã đạt ngay từ "
            "lần đo đầu, chỉ là chọn sai ngưỡng đo lúc đầu. Bài học: định "
            "nghĩa rõ 'đạt' là gì (steady-state, không phải tức thời) trước "
            "khi đo.")
    para(d, "Fail đáng ghi khác: bản apt của gz_ros2_control là cho Gazebo "
            "Fortress, gây segfault ngay khi spawn box trên Harmonic — phải "
            "gỡ và build từ nguồn (~lần đầu tốn thời gian dò nguyên nhân vì "
            "thông báo lỗi không trỏ thẳng tới nguyên nhân phiên bản).")

    h1(d, "M3 — Bắt tay FSM drone<->box qua BoxLink (30/07/2026)")
    para(d, "Đã làm: viết module BoxLink tách riêng khỏi controller; thêm 2 "
            "trạng thái PRELANDING_CHECK/WAIT_BOX_READY; hợp nhất world "
            "Gazebo (marker đặt trên thân box thay vì world riêng); viết "
            "monitor thụ động với 2 tiêu chí nhân quả.")
    para(d, "Insight quan trọng nhất của cả dự án tới nay: một hệ thống có "
            "vẻ 'chạy đúng' (loop khép tới CHARGING) không đủ để kết luận "
            "PASS nếu không chứng minh được tính NHÂN QUẢ — tức là chính sự "
            "kiện A (REQUEST_LANDING) gây ra sự kiện B (box rời EMPTY), chứ "
            "không phải B xảy ra tình cờ do trigger tay còn sót từ lần chạy "
            "trước. Cách chứng minh: đo đúng THỨ TỰ THỜI GIAN giữa hai sự "
            "kiện qua monitor CHỈ NGHE (không publish/gọi service) — nếu "
            "monitor có publish, PASS có thể là do chính monitor kích thích "
            "hệ thống chứ không phải do luồng nghiệp vụ thật. Bài học này "
            "được tái sử dụng nguyên vẹn ở M4 (2 tiêu chí nhân quả tương tự) "
            "và ở kịch bản kiểm tắt-cầu-giữa-chừng.")

    h1(d, "M4 — Tách domain box<->drone qua DDS-Router (05/08/2026)")
    para(d, "Đã làm: khảo sát và build thử DDS-Router 3.5.1 (2 lần độc lập, "
            "2026-07-30 và 2026-08-05) — cả hai lần đều thất bại theo cùng "
            "một cách; chuyển sang DDS-Router 2.2.0, gỡ đúng nguyên nhân gốc "
            "của lỗi bắc cầu ban đầu (whitelist-interfaces per-participant); "
            "chuyển lệnh drone->box từ service sang topic; thêm fixture "
            "/dock/drone_power; viết cầu dự phòng dib_domain_bridge; viết "
            "driver split-domain-aware m4_full_loop_monitor.py; đo bằng "
            "chứng cầu thật qua cổng UDP; kiểm nhân quả tắt cầu giữa chừng "
            "một chuyến bay thật.")
    para(d, "Fail có giá trị định lượng — DDS-Router 3.x: giả thuyết ban đầu "
            "'bản mới nhất chắc tốt hơn' sai hoàn toàn. Build lại lần 2 với "
            "toàn bộ quy trình sạch (không dùng lại kết quả cũ) để loại trừ "
            "khả năng lỗi thao tác — kết quả giống hệt lần 1, xác nhận đây "
            "là giới hạn kiến trúc (Fast DDS 3.x không discovery được "
            "endpoint 2.6), không sửa được bằng cấu hình.")
    para(d, "Fail dẫn tới chẩn đoán sai ban đầu: 'DDS-Router 2.2 cần "
            "ROS_LOCALHOST_ONLY=0' (kết luận 2026-07-30) hoá ra chỉ là "
            "TRIỆU CHỨNG. Bắt gói multicast bằng socket UDP thô mới lộ ra "
            "nguyên nhân thật: whitelist-interfaces là tag per-participant, "
            "đặt nhầm chỗ (top-level) khiến participant router bị node ROS "
            "âm thầm loại bỏ locator. Bài học: khi một 'sửa' có vẻ hiệu quả "
            "nhưng đi kèm đánh đổi lớn (nới lỏng bảo mật mạng), nên nghi ngờ "
            "đó là chẩn đoán sai tầng, không phải nguyên nhân gốc.",)
    para(d, "Fail dẫn tới thiết kế lại giao diện: service b2/cmd không bao "
            "giờ route reply qua domain — xác nhận bằng cả message tuỳ biến "
            "lẫn service ROS 2 chuẩn (loại trừ khả năng lỗi nằm ở dib_msgs). "
            "Quyết định chuyển sang topic không phải là workaround tạm mà "
            "là nhận ra reply của service này chưa từng có ý nghĩa thật từ "
            "M1 (box luôn trả success=true trước khi xử lý xong việc).")
    para(d, "Fail vòng đời (Gazebo lần 1 — Phụ lục E.2 của m4.md): mọi giao "
            "diện hợp đồng qua cầu đúng, hạ cánh chính xác (aim_error=1.2cm), "
            "nhưng box đứng im vĩnh viễn ở SECURING_DRONE, drone đếm hết "
            "90s. Nguyên nhân không phải lỗi cầu mà là DANH SÁCH bắc cầu "
            "thiếu 1 topic (/dock/drone_power) — dấu hiệu chẩn đoán đúng: "
            "thiếu đúng MỘT dòng log kỳ vọng bên phía nhận.")
    para(d, "Kết quả cuối: PASS 8/8 (driver tự động), cộng bằng chứng nhân "
            "quả bậc hai — tắt cầu giữa MỘT CHUYẾN BAY ĐANG SỐNG (khác hẳn "
            "kiểm 'chưa từng bật cầu') cho thấy hệ thống rơi về AUTO.LAND an "
            "toàn đúng thời hạn timeout thiết kế, không treo.")

    h1(d, "M5 — Hardening + debug + đóng gói (đang triển khai)")
    para(d, "Đã làm tới thời điểm viết report này: go/no-go một lệnh "
            "(scripts/go_no_go.sh), heartbeat cho cầu split-domain, cảnh báo "
            "telemetry cũ, mở toàn bộ tài liệu test M1-M4 khỏi .gitignore và "
            "đẩy lên nhánh dev.")
    para(d, "Insight/fail trong quá trình hardening: viết go_no_go.sh lộ ra "
            "một bug thật có sẵn từ trước trong scripts/stop_pipeline.sh — "
            "pattern regex 'spawner' (không có ranh giới từ) khớp nhầm các "
            "tiến trình desktop không liên quan (gvfsd-trash/-network/-dnssd "
            "chạy với cờ '--spawner ...' trên Ubuntu/GNOME). Vì "
            "stop_pipeline.sh dùng kill -9, đây là một bug có khả năng gây "
            "hại thật (kill nhầm tiến trình hệ điều hành), phát hiện được "
            "hoàn toàn tình cờ khi viết công cụ kiểm tra mới dùng lại cùng "
            "pattern — bài học: một script kill -9 với pattern regex rộng "
            "nên được test bằng cách liệt kê PID trước khi tin, không chỉ "
            "dựa vào 'nhìn có vẻ đúng'.")
    para(d, "Fail thứ hai trong go_no_go.sh: bản đầu tiên của mục kiểm "
            "controller-active giả định luôn có 4/4 controller (đúng cho "
            "README mục 2, có box) — chạy thử ở README mục 3 (không box) lộ "
            "ra false-FAIL vì 'ros2 control list_controllers' trả về rỗng "
            "khi không spawn box_simulation. Sửa bằng cách thêm nhánh N_CTRL"
            "==0 là hợp lệ, chỉ FAIL khi có controller nhưng không đủ 4/4 "
            "active.")
    para(d, "Còn lại: hoàn tất bộ tài liệu tái cấu trúc, rà soát go/no-go + "
            "runbook trên máy sạch, chốt trước hạn đóng gói 10/08/2026.")

    d.save(f"{OUT_DIR}/DIB_Report_Giai_Doan.docx")
    print("Đã sinh:", f"{OUT_DIR}/DIB_Report_Giai_Doan.docx")


# ---------------------------------------------------------------------------
# DOC 4 — Hướng dẫn thiết lập/cấu hình + triển khai
# ---------------------------------------------------------------------------

def gen_doc4():
    d = new_doc(
        "HƯỚNG DẪN THIẾT LẬP, CẤU HÌNH & TRIỂN KHAI",
        ["Hệ thống Drone-in-a-Box — SITL-only (không bao gồm chuyển đổi sang "
         "phần cứng thật)",
         "Sinh tự động từ docs/tools/gen_docs.py — nguồn: README.md, "
         "docs/m4_split_domain_test/README.md, verify_build_env.sh, "
         "scripts/go_no_go.sh"],
    )

    h1(d, "1. Yêu cầu môi trường")
    bullets(d, [
        "ROS 2 Humble, source được (/opt/ros/humble/setup.bash).",
        "PX4-Autopilot clone ở ~/PX4, build SITL được (make px4_sitl chạy "
        "được).",
        "Gazebo Harmonic (gz sim --version -> 8.x; đã kiểm trên 8.11.0).",
    ])

    h1(d, "2. Cài đặt một lần — từ máy trống tới build xong")
    para(d, "Toàn bộ 6 bước dưới chỉ cần làm một lần trên một máy.")
    para(d, "Bước 1 — build gz_ros2_control từ nguồn cho Harmonic (bản apt là "
            "cho Fortress, gây segfault):")
    para(d, "sudo apt remove -y ros-humble-gz-ros2-control\n"
            "mkdir -p ~/gz_ros2_control_ws/src && cd ~/gz_ros2_control_ws/src\n"
            "git clone -b humble https://github.com/ros-controls/gz_ros2_control.git\n"
            "cd ~/gz_ros2_control_ws\n"
            "export GZ_VERSION=harmonic\n"
            "rosdep install -r --from-paths src -i -y --rosdistro humble\n"
            "colcon build --symlink-install")
    para(d, "Bước 2 — clone repo vào cây PX4:")
    para(d, "cd ~/PX4/examples\n"
            "git clone <repo-url> SITL_PrecisionLanding\n"
            "cd SITL_PrecisionLanding")
    para(d, "Bước 3 — đồng bộ world/model/texture sang cây PX4:")
    para(d, "cd ~/PX4 && rsync -a "
            "examples/SITL_PrecisionLanding/px4/Tools/simulation/gz/ "
            "Tools/simulation/gz/\n"
            "cd examples/SITL_PrecisionLanding")
    para(d, "Bước 4 — verify (cài dep + build libaruco + giải nén mesh box + "
            "kiểm gz_ros2_control):")
    para(d, "source /opt/ros/humble/setup.bash\n"
            "chmod +x verify_build_env.sh && ./verify_build_env.sh")
    para(d, "Bước 5 — build:")
    para(d, "cd ros2_ws && colcon build --symlink-install && "
            "source install/setup.bash && cd ..")
    para(d, "Bước 6 — dọn tiến trình cũ trước mỗi lần chạy (thói quen, không "
            "chỉ lần đầu):")
    para(d, "./scripts/stop_pipeline.sh   # phải in ra: sach")
    para(d, "box_manager, box_simulation, dib_msgs, precision_landing đã nằm "
            "trong repo — không phải clone thêm gì khác. Bước 1 chỉ cần làm "
            "lại nếu máy chưa từng build gz_ros2_control cho Harmonic.")

    h1(d, "3. Cổng go/no-go — kiểm một lệnh trước mỗi lượt bay")
    para(d, "./scripts/go_no_go.sh")
    para(d, "Tự động hoá 3 việc: môi trường build (verify_build_env.sh), "
            "tiến trình cũ còn sót, và 3 mục kiểm trước bay (use_sim_time, "
            "controller active, MAVROS connected). Tự nhận biết chạy trước "
            "hay sau khi bật pipeline, và tự nhận biết chế độ có/không có "
            "box (README mục 2 và 3) — không báo FAIL sai khi không có "
            "box_simulation. Chỉ đọc (echo/param get/list), không publish, "
            "không gọi service. Chi tiết + checklist mắt người (marker hiện "
            "trên sàn, HUD, arming...): docs/go_no_go.md.")

    h1(d, "4. Chạy chế độ 1 — Drone-in-a-Box đầy đủ (single-domain, 7 "
            "terminal)")
    para(d, "Header cho mọi terminal:")
    para(d, "source /opt/ros/humble/setup.bash\n"
            "source ~/gz_ros2_control_ws/install/setup.bash\n"
            "source ~/PX4/examples/SITL_PrecisionLanding/ros2_ws/install/setup.bash")
    table(d, ["T", "Lệnh", "Vai trò"], [
        ("1", "make px4_sitl gz_x500_gimbal_fractal_aruco_landing (trong ~/PX4)",
         "PX4 + Gazebo, đợi tới pxh>"),
        ("2", "ros2 launch box_simulation box_spawn_only.launch.py",
         "spawn box + marker (đợi ~40s cho 4 controller nạp xong)"),
        ("3", "ros2 launch precision_landing sitl_precland.launch.py",
         "bridge + tracker + controller"),
        ("4", "ros2 launch precision_landing sitl_mavros.launch.py",
         "MAVROS — bắt buộc dùng launch này (đồng hồ mô phỏng)"),
        ("5", "ros2 launch precision_landing dib_bringup.launch.py",
         "cả 3 node phía box"),
        ("6", "ros2 launch .../docs/m3_box_handshake_test/sitl_fixtures.launch.py",
         "fixture GPS, chỉ SITL"),
        ("7", "ros2 run rqt_image_view rqt_image_view /landing/annotated_image",
         "HUD giám sát — mở suốt lượt chạy"),
    ])
    para(d, "Kiểm trước khi bay: ./scripts/go_no_go.sh (mục 3). Bay trong "
            "pxh> của T1: param set NAV_DLL_ACT 0 (nếu không mở "
            "QGroundControl) -> commander takeoff -> commander land. Lượt "
            "chạy đạt = hai FSM đan xen đúng nhân quả, kết thúc ở CHARGING.")

    h1(d, "5. Chạy chế độ 2 — Precision Landing không có box (4 terminal)")
    para(d, "Cùng một binary C++ với chế độ 1. offboard_precland_controller "
            "tự nhận biết không có telemetry box thì bỏ qua nhánh bắt tay, "
            "hạ cánh thị giác tiêu chuẩn.")
    table(d, ["T", "Lệnh", "Vai trò"], [
        ("1", "make px4_sitl gz_x500_gimbal (PX4_GZ_WORLD=fractal_aruco_landing)",
         "PX4 + Gazebo"),
        ("2", "ros2 launch precision_landing sitl_mavros.launch.py", "MAVROS"),
        ("3", "ros2 launch precision_landing sitl_precland.launch.py",
         "bridge + tracker + controller"),
        ("4", "ros2 run rqt_image_view rqt_image_view /landing/annotated_image",
         "HUD"),
    ])
    para(d, "go_no_go.sh chạy được ở đây luôn — tự nhận ra không có "
            "box_simulation (không controller nào) và bỏ qua đúng mục đó "
            "thay vì báo lỗi, chỉ còn kiểm use_sim_time + MAVROS connected.")

    h1(d, "6. Chạy chế độ 3 — Split-domain (box và drone tách "
            "ROS_DOMAIN_ID, chuẩn bị cho hai máy vật lý)")
    para(d, "Prereqs một lần — build DDS-Router 2.2.0 từ nguồn (~4 phút, "
            "KHÔNG có gói apt):")
    para(d, "mkdir -p ~/DDS-Router-2.2/src && cd ~/DDS-Router-2.2\n"
            "cp ~/PX4/examples/SITL_PrecisionLanding/ros2_ws/src/"
            "precision_landing/config/dds_router_2.2.0.repos ddsrouter.repos\n"
            "vcs import src < ddsrouter.repos\n"
            "source /opt/ros/humble/setup.bash && colcon build")
    para(d, "Vì sao 2.2.0 chứ không phải bản mới nhất: mọi DDS-Router 3.x "
            "kéo theo Fast DDS 3.x, không discovery được với Fast DDS 2.6.11 "
            "của Humble. Chi tiết đo đạc: docs/m4.md Phụ lục C.")
    para(d, "Header thêm cho mọi terminal (ngoài 3 dòng source ở mục 2): "
            "export ROS_LOCALHOST_ONLY=1, rồi export ROS_DOMAIN_ID=42 (box) "
            "hoặc =0 (drone) — cầu (T7) không export domain nào.")
    table(d, ["T", "Domain", "Lệnh"], [
        ("1", "42", "PX4 + Gazebo (GZ_SIM_RESOURCE_PATH trỏ vào "
         "install/box_simulation/share)"),
        ("2", "42", "ros2 launch box_simulation box_spawn_only.launch.py"),
        ("5", "42", "ros2 launch precision_landing dib_bringup.launch.py "
         "include_telemetry_bridge:=false"),
        ("6", "42", "ros2 launch .../m3_box_handshake_test/sitl_fixtures.launch.py"),
        ("3", "0", "ros2 launch precision_landing sitl_precland.launch.py"),
        ("4", "0", "ros2 launch precision_landing sitl_mavros.launch.py"),
        ("5b", "0", "ros2 run precision_landing mavros_to_dib_telemetry "
         "--ros-args -p drone_id:=1"),
        ("7", "—", "cầu — DDS-Router (mặc định) hoặc dib_domain_bridge "
         "(dự phòng)"),
    ])
    para(d, "T7 mặc định (DDS-Router):")
    para(d, "cd ~/DDS-Router-2.2 && source install/setup.bash\n"
            "./install/ddsrouter_tool/bin/ddsrouter -c "
            "~/PX4/examples/SITL_PrecisionLanding/ros2_ws/src/"
            "precision_landing/config/dds_router_split.yaml")
    para(d, "T7 dự phòng (domain_bridge, cài từ apt: sudo apt install -y "
            "ros-humble-domain-bridge):")
    para(d, "ros2 run dib_domain_bridge dib_split_bridge 42 0")
    para(d, "Cầu chỉ mang đúng 3 giao diện hợp đồng (b2/telemetry, "
            "d1/telemetry, b2/drone_cmd) + 1 fixture SITL "
            "(/dock/drone_power). Cầu còn sống in heartbeat mỗi 10s "
            "(dib_split_bridge) — không thấy heartbeat mới trong >10s là "
            "cầu chết. Chi tiết run-sheet đầy đủ + cách chứng minh cầu thật "
            "sự mang dữ liệu (không phải domain rò rỉ cùng host): "
            "docs/m4_split_domain_test/README.md.")

    h1(d, "7. Xử lý sự cố")
    table(d, ["Triệu chứng", "Nguyên nhân thật", "Đọc thêm"], [
        ("Máy lag ngay từ đầu lượt chạy mới", "Tiến trình lượt trước còn sót — "
         "chạy ./scripts/stop_pipeline.sh (không gõ tay pkill -f, tự giết "
         "chính shell đang gõ)", "README Phụ lục A.3"),
        ("Controller còn 'unconfigured'", "controller_manager timeout 10s "
         "(Humble) lúc mới khởi tạo — hoàn tất tay 2 bước configure -> "
         "active", "README Phụ lục A.5"),
        ("Tracker vẽ đỏ 'sync N/A: clock mismatch'", "Dùng nhầm 'mavros "
         "px4.launch' thay vì sitl_mavros.launch.py (thiếu use_sim_time)",
         "README Phụ lục A.6"),
        ("(split-domain) box kẹt SECURING_DRONE, không tới CHARGING",
         "Thiếu /dock/drone_power trong allowlist cầu", "docs/m4.md Phụ lục E.2"),
        ("(split-domain) box kẹt EMPTY, WAIT_BOX_READY timeout 30s -> "
         "FALLBACK", "Cầu chết hoặc thiếu whitelist-interfaces per-participant",
         "docs/m4.md Phụ lục B.1, I.2"),
        ("Bảng đầy đủ hơn (12 dòng, gồm cả mắt-người-cần-kiểm)", "—",
         "docs/go_no_go.md"),
    ])

    d.save(f"{OUT_DIR}/DIB_Huong_Dan_Thiet_Lap_Trien_Khai.docx")
    print("Đã sinh:", f"{OUT_DIR}/DIB_Huong_Dan_Thiet_Lap_Trien_Khai.docx")


if __name__ == "__main__":
    gen_doc1()
    gen_doc2()
    gen_doc3()
    gen_doc4()
